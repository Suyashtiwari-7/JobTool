"""Resume parser — extracts text from PDF/DOCX and structures it via LLM."""

import json
import logging

try:
    import fitz  # PyMuPDF
except ImportError:
    try:
        import pymupdf as fitz
    except ImportError:
        fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

from app.llm.provider import llm_call
from app.llm.prompts import RESUME_PARSE_PROMPT

logger = logging.getLogger(__name__)


COMMON_SKILLS = [
    "Python", "JavaScript", "TypeScript", "React", "Next.js", "Node.js", "FastAPI",
    "Express", "Django", "Flask", "HTML", "CSS", "SQL", "PostgreSQL", "MongoDB",
    "Docker", "Kubernetes", "AWS", "GCP", "Azure", "Git", "GitHub", "REST API",
    "GraphQL", "C++", "Java", "C#", "Go", "Rust", "Tailwind", "Machine Learning",
    "AI", "PyTorch", "TensorFlow", "Pandas", "NumPy"
]


def _fast_parse_text(raw_text: str) -> dict:
    """Instant deterministic extraction of contact info and skills via regex."""
    import re
    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
    name = lines[0] if lines else "Candidate"

    # Extract Email
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', raw_text)
    email = email_match.group(0) if email_match else None

    # Extract Phone
    phone_match = re.search(r'(\+?\d{1,3}[\s-]?)?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}', raw_text)
    phone = phone_match.group(0) if phone_match else None

    # Extract Skills
    found_skills = [
        skill for skill in COMMON_SKILLS
        if re.search(r'\b' + re.escape(skill) + r'\b', raw_text, re.IGNORECASE)
    ]

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "location": None,
        "summary": lines[1] if len(lines) > 1 else raw_text[:250],
        "skills": found_skills,
        "experience": [],
        "education": [],
        "certifications": [],
        "links": []
    }


async def parse_resume_file(file_path: str, extension: str) -> tuple[str, dict]:
    """
    Parse a resume file into structured JSON.
    Instant local extraction. LLM tailoring & scoring happens later during job application pipeline.
    """
    # Step 1: Extract raw text
    if extension == ".pdf":
        raw_text = _extract_pdf_text(file_path)
    elif extension == ".docx":
        raw_text = _extract_docx_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")

    if not raw_text.strip():
        raw_text = "Uploaded Resume Document"

    logger.info(f"Extracted {len(raw_text)} chars from resume")

    # Step 2: Instant deterministic extraction
    parsed = _fast_parse_text(raw_text)

    logger.info(f"Parsed resume locally for: {parsed.get('name', 'Unknown')}")
    return raw_text, parsed


def _extract_pdf_text(file_path: str) -> str:
    """Extract text from a PDF using PyMuPDF."""
    if fitz is not None:
        try:
            doc = fitz.open(file_path)
            text_parts: list[str] = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            text = "\n".join(text_parts).strip()
            if text:
                return text
        except Exception as e:
            logger.warning(f"PDF extraction error ({e}). Returning fallback text.")

    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return "Uploaded Resume Document"


def _extract_docx_text(file_path: str) -> str:
    """Extract text from a DOCX using python-docx."""
    if Document is not None:
        try:
            doc = Document(file_path)
            text_parts: list[str] = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            text = "\n".join(text_parts).strip()
            if text:
                return text
        except Exception as e:
            logger.warning(f"DOCX extraction error ({e}). Returning fallback text.")

    return "Uploaded Resume Document"
